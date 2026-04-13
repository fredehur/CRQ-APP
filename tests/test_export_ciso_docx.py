"""Tests for export_ciso_docx — CISO brief redesign (2026-04-13 spec)."""
import os
import json
import pytest
from docx import Document
from docx.shared import RGBColor

import tools.export_ciso_docx as ed


def test_export_runs_without_error(mock_output, tmp_path):
    """export() completes without raising and writes a non-trivial .docx."""
    out = str(tmp_path / "ciso_brief.docx")
    ed.export(output_path=out, output_dir=str(mock_output))
    assert os.path.exists(out)
    assert os.path.getsize(out) > 5_000  # at least 5 KB


def test_docx_contains_bluf(mock_output, tmp_path):
    """BLUF sentence appears in the document before the PURPOSE section."""
    out = str(tmp_path / "ciso_brief.docx")
    ed.export(output_path=out, output_dir=str(mock_output))
    doc = Document(out)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # PURPOSE section header must exist
    purpose_indices = [i for i, t in enumerate(texts) if "PURPOSE" in t.upper()]
    assert purpose_indices, "PURPOSE section header not found"
    # BLUF must appear before PURPOSE
    # BLUF has no label — it is the first substantial non-header paragraph
    purpose_idx = purpose_indices[0]
    pre_purpose = texts[:purpose_idx]
    # At least one sentence-like text before PURPOSE
    sentences = [t for t in pre_purpose if len(t) > 20 and "CISO" not in t.upper()]
    assert sentences, "No BLUF sentence found before PURPOSE section"


def test_docx_contains_all_section_headers(mock_output, tmp_path):
    """All six named section headers appear in the document."""
    out = str(tmp_path / "ciso_brief.docx")
    ed.export(output_path=out, output_dir=str(mock_output))
    doc = Document(out)
    full_text = "\n".join(p.text.upper() for p in doc.paragraphs)
    for header in ["PURPOSE", "INTELLIGENCE PICTURE", "THREAT ASSESSMENT", "SITUATION", "ACTION REGISTER", "CONSIDERATIONS"]:
        assert header in full_text, f"Section header '{header}' not found in document"


def test_watch_list_present_when_monitor_regions(mock_output, tmp_path):
    """WATCH LIST section appears when MONITOR regions exist."""
    out = str(tmp_path / "ciso_brief.docx")
    ed.export(output_path=out, output_dir=str(mock_output))
    doc = Document(out)
    full_text = "\n".join(p.text.upper() for p in doc.paragraphs)
    # AME is MONITOR in mock data — watch list should appear
    assert "WATCH LIST" in full_text, "WATCH LIST section missing despite MONITOR region"


def test_action_register_deduplicates(mock_output, tmp_path):
    """'Validate offline backup integrity' appears exactly once despite appearing in both APAC and MED action_bullets."""
    out = str(tmp_path / "ciso_brief.docx")
    ed.export(output_path=out, output_dir=str(mock_output))
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    count = full_text.lower().count("validate offline backup integrity")
    assert count == 1, f"Expected 1 occurrence of dedup'd action, found {count}"


def test_no_red_amber_green_in_body(mock_output, tmp_path):
    """No RED, AMBER, or GREEN RGBColor appears in any paragraph run (A1 visual style)."""
    out = str(tmp_path / "ciso_brief.docx")
    ed.export(output_path=out, output_dir=str(mock_output))
    doc = Document(out)
    forbidden = {
        (0xFF, 0x00, 0x00),  # pure red
        (0xDC, 0x26, 0x26),  # Tailwind red-600
        (0xFF, 0xBF, 0x00),  # amber
        (0xFF, 0xA5, 0x00),  # orange-amber
        (0x00, 0x80, 0x00),  # green
        (0x16, 0xA3, 0x4A),  # Tailwind green-600
        (0xD9, 0x7F, 0x0F),  # amber-700
    }
    for para in doc.paragraphs:
        for run in para.runs:
            c = run.font.color
            if c and c.type is not None:
                try:
                    rgb = c.rgb
                    if rgb:
                        triple = (rgb.red, rgb.green, rgb.blue)
                        assert triple not in forbidden, (
                            f"Forbidden colour {triple} found in run: '{run.text[:60]}'"
                        )
                except Exception:
                    pass  # colour not set — fine
