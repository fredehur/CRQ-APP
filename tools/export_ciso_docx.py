"""
tools/export_ciso_docx.py
Exports a CISO-grade weekly intelligence brief as a Word (.docx) document.

Redesigned 2026-04-13 per docs/superpowers/specs/2026-04-13-ciso-brief-redesign.md
Visual style: A1 (pure type, black/white/grey only — no RED/AMBER/GREEN).
Structure: BLUF → Purpose → Intelligence Picture → Threat Assessment → Situation →
           Watch List → Action Register → Considerations → References → Footer.

Usage:
    uv run python tools/export_ciso_docx.py [output.docx]
    Defaults to output/ciso_brief.docx
"""
import sys
import os
import re
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(__file__))
sys.path.insert(0, str(Path(__file__).parent.parent))
from report_builder import build, ReportData, RegionEntry, RegionStatus
from tools.config import CISO_BRIEF_PATH

DEFAULT_OUT = str(CISO_BRIEF_PATH)
OUTPUT_DIR  = "output/pipeline"

# ── Colours (A1 style: black/white/grey only) ────────────────────────────────
GREY  = RGBColor(0x6B, 0x72, 0x80)
BODY  = RGBColor(0x11, 0x11, 0x11)
BLACK = BODY  # alias — kept for clarity in helpers

# ── Region display names ───────────────────────────────────────────────────────
REGION_NAMES = {
    "AME":   "Americas, Europe & Middle East",
    "APAC":  "Asia-Pacific",
    "MED":   "Mediterranean",
    "NCE":   "Northern & Central Europe",
    "LATAM": "Latin America",
}

# ── Credibility tier inference ────────────────────────────────────────────────
_TIER_A_DOMAINS = (".gov", "enisa.europa.eu", "nist.gov", "cisa.gov",
                   "eur-lex.europa.eu", "nato.int", "un.org", "interpol.int")
_TIER_C_DOMAINS = ("youtube.com", "twitter.com", "reddit.com", "facebook.com")


def _infer_tier(url: str | None) -> str:
    """Return credibility tier A/B/C based on URL domain patterns."""
    if not url:
        return "B"
    low = url.lower()
    for domain in _TIER_A_DOMAINS:
        if domain in low:
            return "A"
    for domain in _TIER_C_DOMAINS:
        if domain in low:
            return "C"
    return "B"


# Sources to skip when building the reference list (generic collector labels)
_SKIP_SOURCES = {
    "tavily research", "cyber signal", "geo signal", "youtube signal",
    "research collection", "osint collection",
}
_SKIP_PREFIXES = ("cyber signal", "geo signal", "youtube signal", "tavily")

# Bracket patterns that are NOT citations
_NON_CITATION_RE = re.compile(r"^(ESCALATED|MONITOR|CLEAR|UNKNOWN)$")

# Sentence splitter (used by _split_why and watch list first-sentence extraction)
_SENT_RE = re.compile(r'(?<=[.!?])\s+(?=[A-Z"\'(])')


# ── Source registry ────────────────────────────────────────────────────────────
class SourceRegistry:
    """Global numbered reference registry for in-text citations."""

    def __init__(self, url_map: dict[str, str] | None = None):
        self._refs: dict[str, dict] = {}
        self._counter = 0
        # name.lower() → url, populated from signal source files
        self._url_map: dict[str, str] = url_map or {}

    def _key(self, name: str) -> str:
        return name.strip().lower()

    def register(self, name: str, headline: str | None = None) -> int:
        k = self._key(name)
        if k not in self._refs:
            self._counter += 1
            self._refs[k] = {
                "num": self._counter,
                "name": name.strip(),
                "headline": headline or "",
                "url": self._url_map.get(k, ""),
            }
        else:
            if headline and not self._refs[k]["headline"]:
                self._refs[k]["headline"] = headline
            if not self._refs[k].get("url"):
                self._refs[k]["url"] = self._url_map.get(k, "")
        return self._refs[k]["num"]

    def lookup(self, name: str) -> int | None:
        return self._refs.get(self._key(name), {}).get("num")

    def all_refs(self) -> list[dict]:
        return sorted(self._refs.values(), key=lambda r: r["num"])


# ── Load signal sources (geo + cyber) → name → url map ────────────────────────
def _load_signal_sources(region_lower: str, output_dir: str = OUTPUT_DIR) -> dict[str, str]:
    """Read osint_signals.json for a region.

    Returns a dict mapping source name.lower() → url.
    Returns empty dict if file is absent or the `sources` field is missing.
    """
    url_map: dict[str, str] = {}
    for filename in ("osint_signals.json",):
        path = Path(output_dir).parent / "regional" / region_lower / filename
        if not path.exists():
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        for src in data.get("sources", []):
            name = src.get("name", "").strip()
            url  = src.get("url", "").strip()
            if name and url:
                url_map[name.lower()] = url
    return url_map


def _build_signal_url_map(output_dir: str = OUTPUT_DIR) -> dict[str, str]:
    """Merge signal source URL maps across all regions."""
    merged: dict[str, str] = {}
    for region in ["apac", "ame", "med", "nce", "latam"]:
        for k, v in _load_signal_sources(region, output_dir).items():
            if k not in merged:
                merged[k] = v
    return merged


# ── Load source clusters ───────────────────────────────────────────────────────
def _load_clusters(region_lower: str, output_dir: str = OUTPUT_DIR) -> dict[str, str]:
    path = Path(output_dir).parent / "regional" / region_lower / "signal_clusters.json"
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    result: dict[str, str] = {}
    for cluster in data.get("clusters", []):
        for src in cluster.get("sources", []):
            name     = src.get("name", "").strip()
            headline = src.get("headline", "").strip()
            if name and name.lower() not in _SKIP_SOURCES:
                result[name.lower()] = headline
    return result


def _build_cluster_map(output_dir: str = OUTPUT_DIR) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for region in ["apac", "ame", "med", "nce", "latam"]:
        for k, v in _load_clusters(region, output_dir).items():
            if k not in mapping:
                mapping[k] = v
    return mapping


# ── Citation processing ────────────────────────────────────────────────────────
_BRACKET_RE = re.compile(r"\[([^\]]+)\]")
# Matches "Evidenced by [N]: " or "Corroborated by [N]: " prefixes
_ATTRIBUTION_RE = re.compile(
    r"(?:Evidenced by|Corroborated by|Source:|Cited by|Confirmed by)"
    r"\s+(\[\d+(?:,\s*\d+)*\]):\s*",
    re.IGNORECASE,
)


def _process_citations(text: str, registry: SourceRegistry,
                        cluster_map: dict[str, str]) -> str:
    """Replace [Source Name] with [N] and clean attribution prefixes."""
    def _register_one(name: str) -> int:
        return registry.register(name, cluster_map.get(name.lower(), ""))

    def replace(match: re.Match) -> str:
        raw = match.group(1).strip()
        if _NON_CITATION_RE.match(raw):
            return match.group(0)
        if raw.lower() in _SKIP_SOURCES:
            return match.group(0)
        if any(raw.lower().startswith(p) for p in _SKIP_PREFIXES):
            return match.group(0)
        parts = [p.strip() for p in raw.split(",") if p.strip()]
        if len(parts) > 1:
            return "[" + ", ".join(str(_register_one(p)) for p in parts) + "]"
        return f"[{_register_one(raw)}]"

    # Step 1: replace source names with numbers
    text = _BRACKET_RE.sub(replace, text)

    # Step 2: strip "Evidenced by [N]: " / "Corroborated by [N]: " prefixes
    def move_citation(m: re.Match) -> str:
        return ""

    text = _ATTRIBUTION_RE.sub(move_citation, text)
    return text


# ── Document helpers ───────────────────────────────────────────────────────────
def _font(run, size: int, bold: bool = False, italic: bool = False,
          colour: RGBColor | None = None) -> None:
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour


def _add_rule(doc: Document) -> None:
    """Insert a thin horizontal divider (grey bottom border on empty paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D1D5DB")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_subheading(doc: Document, text: str) -> None:
    """A1-style subheading: small grey uppercase with a thin rule below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    _font(run, 8, bold=True, colour=GREY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '9CA3AF')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


# ── Data access helpers (sections.json, runs log) ─────────────────────────────
def _split_why(text: str, max_bullets: int = 3) -> list[str]:
    """Split a paragraph into 2-3 sentence bullets for talking points."""
    if not text or not text.strip():
        return []
    sentences = [s.strip() for s in _SENT_RE.split(text.strip()) if s.strip()]
    return sentences[:max_bullets]


def _group_by_scenario(entries: list[RegionEntry]) -> list[tuple[str, list[RegionEntry]]]:
    """Group RegionEntry list by scenario_match. Returns [(scenario, [entries])]."""
    groups: dict[str, list[RegionEntry]] = {}
    for entry in entries:
        key = getattr(entry, 'scenario_match', None) or f"Unknown — {entry.name}"
        groups.setdefault(key, []).append(entry)
    return list(groups.items())


def _load_sections(region_name: str, output_dir: str = OUTPUT_DIR) -> dict:
    """Load sections.json for a region. Regional tree is sibling of pipeline dir."""
    path = Path(output_dir).parent / "regional" / region_name.lower() / "sections.json"
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _load_source_signal(region_name: str, output_dir: str = OUTPUT_DIR) -> str:
    """Return '(N sources · N corroborated)' string, or '' if data missing."""
    sec = _load_sections(region_name, output_dir)
    osint = sec.get("source_metadata", {}).get("osint", {})
    count = osint.get("source_count")
    tier = osint.get("corroboration_tier")
    if count is None:
        return ""
    if tier:
        return f"({count} sources · {tier})"
    return f"({count} sources)"


def _load_brief_headlines(region_name: str, output_dir: str = OUTPUT_DIR) -> dict:
    """Load brief_headlines from sections.json for a region."""
    sec = _load_sections(region_name, output_dir)
    return sec.get("brief_headlines", {})


def _get_previous_escalated_count(output_dir: str = OUTPUT_DIR) -> int | None:
    """Read last archived run log and return escalated count, or None."""
    runs_dir = Path(output_dir) / "runs"
    if not runs_dir.is_dir():
        return None
    try:
        logs = sorted(
            [f for f in os.listdir(runs_dir) if f.endswith(".json")],
            reverse=True
        )
        if len(logs) < 2:
            return None
        prev_path = runs_dir / logs[1]  # second-most-recent
        data = json.loads(prev_path.read_text(encoding="utf-8"))
        regions = data.get("regions", {})
        return sum(1 for r in regions.values() if r.get("status") == "ESCALATED")
    except Exception:
        return None


# ── Derived report metadata ────────────────────────────────────────────────────
def _report_date(data: ReportData) -> str:
    return (data.timestamp[:10]
            if data.timestamp and data.timestamp != "unknown"
            else datetime.now().strftime("%Y-%m-%d"))


def _cycle_label(data: ReportData) -> str:
    return data.run_id or "N/A"


# ── Section builders ───────────────────────────────────────────────────────────
def _build_cover(doc: Document, data: ReportData) -> None:
    """Minimal A1 cover — title, date/cycle line, thin divider. No page break."""
    p = doc.add_paragraph()
    run = p.add_run("CISO Intelligence Brief")
    _font(run, 15, bold=True, colour=BODY)

    p2 = doc.add_paragraph()
    meta = f"{_report_date(data)}  ·  Cycle {_cycle_label(data)}"
    run2 = p2.add_run(meta)
    _font(run2, 9, colour=GREY)

    _add_rule(doc)


def _build_bluf(doc: Document, data: ReportData) -> None:
    """One-sentence Bottom Line Up Front — no label, bold, top of doc."""
    escalated = [e for e in data.regions if e.status == RegionStatus.ESCALATED]
    if escalated:
        names = ", ".join(e.name for e in escalated)
        scenarios = list({
            getattr(e, 'scenario_match', None) or 'unknown threat'
            for e in escalated
        })
        scenario_str = " and ".join(scenarios[:2])
        if len(escalated) >= 2:
            bluf = (
                f"Two threats require immediate attention this week: "
                f"{scenario_str} affecting {names}."
            )
        else:
            bluf = (
                f"One threat requires immediate attention this week: "
                f"{scenario_str} — {names}."
            )
    else:
        monitor_count = sum(1 for e in data.regions if e.status == RegionStatus.MONITOR)
        plural = 's' if monitor_count != 1 else ''
        bluf = (
            f"No threats escalated this cycle — "
            f"{monitor_count} region{plural} remain on monitor."
        )
    p = doc.add_paragraph()
    run = p.add_run(bluf)
    _font(run, 11, bold=True, colour=BODY)
    p.paragraph_format.space_after = Pt(12)


def _build_purpose(doc: Document, data: ReportData) -> None:
    """Static boilerplate with dynamic date/cycle/region count."""
    _add_subheading(doc, "Purpose of This Brief")
    region_count = len(data.regions)
    region_names = ", ".join(e.name for e in data.regions)
    report_date = _report_date(data)
    text = (
        f"This brief provides AeroGrid Wind Solutions' CISO with a consolidated geopolitical "
        f"and cyber threat assessment for the week of {report_date}. Intelligence current as of "
        f"{report_date}. It covers {region_count} operational regions ({region_names}) and is produced "
        f"by the CRQ agentic intelligence pipeline (Cycle {_cycle_label(data)}). Intended for internal "
        f"decision-making and upward briefing. Not for external distribution."
    )
    p = doc.add_paragraph()
    run = p.add_run(text)
    _font(run, 10, colour=BODY)


def _build_intelligence_picture(doc: Document, data: ReportData) -> None:
    """Status summary line + cycle delta + global framing paragraph."""
    _add_subheading(doc, "Intelligence Picture")

    escalated = sum(1 for e in data.regions if e.status == RegionStatus.ESCALATED)
    monitored = sum(1 for e in data.regions if e.status == RegionStatus.MONITOR)
    cleared   = sum(1 for e in data.regions if e.status == RegionStatus.CLEAR)

    prev = _get_previous_escalated_count(getattr(data, 'output_dir', OUTPUT_DIR))
    if prev is None:
        delta_str = ""
    elif prev == escalated:
        delta_str = " (unchanged)"
    else:
        diff = escalated - prev
        delta_str = f" ({'+' if diff > 0 else ''}{diff} since previous cycle)"

    p = doc.add_paragraph()
    run = p.add_run(
        f"{escalated} escalated  ·  {monitored} monitored  ·  {cleared} clear{delta_str}"
    )
    _font(run, 10, bold=True, colour=BODY)

    # Global framing
    summary = getattr(data, 'exec_summary', '') or ''
    if not summary:
        esc_names = ", ".join(e.name for e in data.regions if e.status == RegionStatus.ESCALATED)
        pillars = list({
            (getattr(e, 'dominant_pillar', None) or 'CYBER').upper()
            for e in data.regions
            if e.status == RegionStatus.ESCALATED
        })
        dominant = pillars[0] if pillars else "CYBER"
        if esc_names:
            plural = 'are' if escalated > 1 else 'is'
            mon_plural = 's' if monitored != 1 else ''
            summary = (
                f"This cycle, {esc_names} {plural} escalated. "
                f"The dominant threat type is {dominant}-led. "
                f"{monitored} region{mon_plural} remain on monitor."
            )
        else:
            mon_plural = 's' if monitored != 1 else ''
            summary = (
                f"No threats escalated this cycle. "
                f"{monitored} region{mon_plural} remain on monitor."
            )

    p2 = doc.add_paragraph()
    run2 = p2.add_run(summary)
    _font(run2, 10, colour=BODY)


def _build_threat_assessments(doc: Document, data: ReportData,
                                registry: SourceRegistry,
                                cluster_map: dict[str, str]) -> None:
    """Section 2: one block per threat scenario, regions grouped."""
    escalated = [e for e in data.regions if e.status == RegionStatus.ESCALATED]
    if not escalated:
        return

    _add_subheading(doc, "Threat Assessment")

    output_dir = getattr(data, 'output_dir', OUTPUT_DIR)

    for scenario, group in _group_by_scenario(escalated):
        region_names = ", ".join(e.name for e in group)

        # Scenario header
        p = doc.add_paragraph()
        run = p.add_run(f"{scenario}  —  {region_names}")
        _font(run, 12, bold=True, colour=BODY)

        # ESCALATED badge + pillar + admiralty (simulated inline — docx runs have
        # no native background colour support, so we use bold uppercase + grey meta)
        rep = group[0]  # representative entry
        p2 = doc.add_paragraph()
        badge = p2.add_run("ESCALATED")
        _font(badge, 8, bold=True, colour=BODY)

        pillar = (getattr(rep, 'dominant_pillar', None) or 'CYBER').upper()
        admiralty = getattr(rep, 'admiralty', '') or ''
        meta_run = p2.add_run(f"  {pillar}-LED  ·  Admiralty {admiralty}")
        _font(meta_run, 8, colour=GREY)

        # Source signal line (from first region in group)
        sig = _load_source_signal(group[0].name, output_dir)
        if sig:
            p3 = doc.add_paragraph()
            sig_run = p3.add_run(sig)
            _font(sig_run, 8, colour=GREY)

        # Talking points — from brief_headlines.why, fall back to why_text
        bh = _load_brief_headlines(group[0].name, output_dir)
        why = bh.get("why") or getattr(rep, 'why_text', '') or ''
        bullets = _split_why(why)
        if bullets:
            admcode = f" [{admiralty}]" if admiralty else ""
            for bullet in bullets:
                p_b = doc.add_paragraph(style='List Bullet')
                run_b = p_b.add_run(f"{bullet}{admcode}")
                _font(run_b, 10, colour=BODY)

        # Impact — from brief_headlines.so_what, fall back to so_what_text
        so_what = bh.get("so_what") or getattr(rep, 'so_what_text', '') or ''
        if so_what:
            p_sw = doc.add_paragraph()
            label = p_sw.add_run("Impact: ")
            _font(label, 10, bold=True, colour=BODY)
            val = p_sw.add_run(so_what)
            _font(val, 10, colour=BODY)

        doc.add_paragraph()  # spacer


def _build_situation(doc: Document, data: ReportData,
                       registry: SourceRegistry,
                       cluster_map: dict[str, str]) -> None:
    """Section 3: one block per threat scenario — what happened and how."""
    escalated = [e for e in data.regions if e.status == RegionStatus.ESCALATED]
    if not escalated:
        return

    _add_subheading(doc, "Situation")

    output_dir = getattr(data, 'output_dir', OUTPUT_DIR)

    for scenario, group in _group_by_scenario(escalated):
        rep = group[0]
        region_names = ", ".join(e.name for e in group)

        # Header
        p = doc.add_paragraph()
        run = p.add_run(f"{scenario}  —  {region_names}")
        _font(run, 11, bold=True, colour=BODY)

        # Narrative — brief_headlines.how, fallback how_text
        bh = _load_brief_headlines(rep.name, output_dir)
        how = bh.get("how") or getattr(rep, 'how_text', '') or ''
        if how:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(how)
            _font(run2, 10, colour=BODY)

        # Supporting intel bullets (max 3)
        sec = _load_sections(rep.name, output_dir)
        intel_bullets = sec.get("intel_bullets", [])[:3]
        for bullet in intel_bullets:
            p_b = doc.add_paragraph(style='List Bullet')
            run_b = p_b.add_run(bullet)
            _font(run_b, 10, colour=BODY)

        doc.add_paragraph()  # spacer


def _build_watch_list(doc: Document, data: ReportData) -> None:
    """Section 4: one line per MONITOR region. Omit section if none."""
    monitor = [e for e in data.regions if e.status == RegionStatus.MONITOR]
    if not monitor:
        return

    _add_subheading(doc, "Watch List")

    output_dir = getattr(data, 'output_dir', OUTPUT_DIR)

    for entry in monitor:
        bh = _load_brief_headlines(entry.name, output_dir)
        why = bh.get("why") or getattr(entry, 'why_text', '') or ''
        first_sentence = (
            _SENT_RE.split(why)[0].strip()
            if why
            else f"{entry.name} — on monitor."
        )
        p = doc.add_paragraph()
        region_run = p.add_run(f"{entry.name} — ")
        _font(region_run, 10, bold=True, colour=BODY)
        text_run = p.add_run(first_sentence)
        _font(text_run, 10, colour=BODY)


def _build_action_register(doc: Document, data: ReportData) -> None:
    """Section 5: deduplicated actions from all escalated regions."""
    _add_subheading(doc, "Action Register")

    escalated = [e for e in data.regions if e.status == RegionStatus.ESCALATED]
    output_dir = getattr(data, 'output_dir', OUTPUT_DIR)

    # action_map: {action_text: [region_names]}
    action_map: dict[str, list[str]] = {}
    for entry in escalated:
        sec = _load_sections(entry.name, output_dir)
        for bullet in sec.get("action_bullets", []):
            key = bullet.strip()
            if not key:
                continue
            action_map.setdefault(key, [])
            if entry.name not in action_map[key]:
                action_map[key].append(entry.name)

    if not action_map:
        p = doc.add_paragraph()
        run = p.add_run("No specific actions required this cycle.")
        _font(run, 10, colour=BODY)
        return

    # Sort: actions tagged across more regions first, then longer (more specific) first
    sorted_actions = sorted(
        action_map.items(),
        key=lambda x: (-len(x[1]), -len(x[0])),
    )

    for i, (action, regions) in enumerate(sorted_actions, 1):
        region_tag = f" — {', '.join(regions)}" if regions else ""
        p = doc.add_paragraph()
        num_run = p.add_run(f"{i}. ")
        _font(num_run, 10, bold=True, colour=BODY)
        text_run = p.add_run(f"{action}{region_tag}")
        _font(text_run, 10, colour=BODY)


def _build_considerations(doc: Document, data: ReportData,
                            registry: SourceRegistry,
                            cluster_map: dict[str, str]) -> None:
    """Section 6: decision-relevant framing — watch bullets + global framing."""
    _add_subheading(doc, "Considerations")

    output_dir = getattr(data, 'output_dir', OUTPUT_DIR)
    bullets_added = 0

    # From watch_bullets across all escalated regions
    escalated = [e for e in data.regions if e.status == RegionStatus.ESCALATED]
    seen: set[str] = set()
    for entry in escalated:
        sec = _load_sections(entry.name, output_dir)
        for wb in sec.get("watch_bullets", []):
            if wb not in seen and bullets_added < 4:
                seen.add(wb)
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(wb)
                _font(run, 10, colour=BODY)
                bullets_added += 1

    # From exec_summary (global layer)
    summary = getattr(data, 'exec_summary', '') or ''
    if summary and bullets_added < 4:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(summary[:300])
        _font(run, 10, colour=BODY)
        bullets_added += 1

    if bullets_added == 0:
        p = doc.add_paragraph()
        run = p.add_run("No strategic considerations identified this cycle.")
        _font(run, 10, colour=BODY)


def _build_references(doc: Document, registry: SourceRegistry) -> None:
    refs = registry.all_refs()
    if not refs:
        return

    _add_subheading(doc, "References")

    # Sort by credibility tier: A first, then B, then C — preserve order within tier
    _TIER_ORDER = {"A": 0, "B": 1, "C": 2}
    sorted_refs = sorted(
        refs,
        key=lambda r: _TIER_ORDER.get(_infer_tier(r.get("url", "")), 1),
    )

    for ref in sorted_refs:
        name = ref["name"]
        url  = ref.get("url", "")
        tier = _infer_tier(url)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent       = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.space_after       = Pt(4)

        r1 = p.add_run(f"[{tier}] ")
        _font(r1, 9, bold=True, colour=BODY)

        body = f"{name} \u2014 {url}" if url else name
        r2 = p.add_run(body)
        _font(r2, 9, colour=BODY)


def _build_footer_note(doc: Document, data: ReportData) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        f"Pipeline run: {data.run_id}  |  "
        "Sources: open-source signals (media monitoring, government publications, industry reporting).  |  "
        "Scenario register: AeroGrid CRQ database.  |  "
        "For internal use only."
    )
    _font(r, 8, italic=True, colour=GREY)


# ── Main ───────────────────────────────────────────────────────────────────────
def export(output_path: str = DEFAULT_OUT, output_dir: str = OUTPUT_DIR) -> None:
    """Generate CISO Intelligence Brief as .docx."""
    data = build(output_dir=output_dir)
    # Expose output_dir on the data object so builder helpers can resolve the
    # regional/ tree (sibling of the pipeline dir).
    setattr(data, 'output_dir', output_dir)

    url_map     = _build_signal_url_map(output_dir)
    registry    = SourceRegistry(url_map=url_map)
    cluster_map = _build_cluster_map(output_dir)

    doc = Document()
    _set_margins(doc)

    _build_cover(doc, data)
    _build_bluf(doc, data)
    _build_purpose(doc, data)
    _build_intelligence_picture(doc, data)
    _build_threat_assessments(doc, data, registry, cluster_map)
    _build_situation(doc, data, registry, cluster_map)
    _build_watch_list(doc, data)
    _build_action_register(doc, data)
    _build_considerations(doc, data, registry, cluster_map)
    _build_references(doc, registry)
    _build_footer_note(doc, data)

    doc.save(output_path)
    print(f"[export_ciso_docx] Saved -> {output_path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_OUT
    export(out)
